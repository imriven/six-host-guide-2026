from pathlib import Path

from xlsx_to_json import load_workbook_data

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
_, generated_run_of_show = load_workbook_data()
rows = [
    {
        "kind": str(item["kind"]).title(),
        "segment": str(item["segmentNumber"] or ""),
        "start": str(item["start"]),
        "end": str(item["end"]),
        "hosts": str(item["hosts"]),
        "title": str(item["title"]),
        "developer": str(item["developer"]),
        "notes": str(item["productionNotes"]),
    }
    for item in generated_run_of_show["items"]
]
output = ROOT / "public" / "SIX-2026-Run-of-Show.docx"

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def write_cell(cell, text, *, bold=False, italic=False, color="202028", size=7.2):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width, section.page_height = Inches(11), Inches(8.5)
section.top_margin = section.bottom_margin = Inches(.42)
section.left_margin = section.right_margin = Inches(.42)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(1)
run = title.add_run("Seattle Indies Expo - SIX 2026")
run.font.name = "Arial"; run.font.size = Pt(19); run.bold = True; run.font.color.rgb = RGBColor(116, 69, 199)
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(4)
run = subtitle.add_run("Run of Show - Rough Draft")
run.font.name = "Arial"; run.font.size = Pt(10); run.bold = True; run.font.color.rgb = RGBColor(214, 69, 36)
hosts = doc.add_paragraph(str(generated_run_of_show["hostNote"]).replace(" | ", "   |   "))
hosts.paragraph_format.space_after = Pt(7)
hosts.runs[0].font.name = "Arial"; hosts.runs[0].font.size = Pt(8)

headers = ["SEG", "GAME / BREAK", "DEVELOPER / PUBLISHER", "START", "END", "HOSTS", "PRODUCTION NOTES"]
widths = [.38, 1.72, 1.45, .70, .70, .62, 4.76]
table = doc.add_table(rows=1, cols=len(headers))
table.autofit = False
for index, header in enumerate(headers):
    table.columns[index].width = Inches(widths[index])
    write_cell(table.rows[0].cells[index], header, bold=True, color="FFFFFF", size=7.5)
    shade(table.rows[0].cells[index], "7445C7")
table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))

for row in rows:
    cells = table.add_row().cells
    values = [row["segment"] or "-", row["title"], row["developer"], row["start"], row["end"], row["hosts"], row["notes"]]
    for index, value in enumerate(values):
        is_title = index == 1
        write_cell(cells[index], value,
                   bold=is_title and row["kind"] in ("Game", "Intro", "Closing", "Ad"),
                   italic=is_title and row["kind"] == "Transition",
                   color="7445C7" if is_title and row["kind"] == "Game" else "202028")
        cells[index].width = Inches(widths[index])
    if row["kind"] == "Ad":
        for cell in cells: shade(cell, "EADCF4")
    elif row["kind"] == "Transition":
        for cell in cells: shade(cell, "F2F1F5")
    elif row["kind"] in ("Intro", "Closing"):
        for cell in cells: shade(cell, "E4F6D9")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = footer.add_run("SIX 2026 Run of Show - Rough Draft")
run.font.name = "Arial"; run.font.size = Pt(7); run.font.color.rgb = RGBColor(102, 97, 110)
doc.save(output)
print({"rows": len(rows), "output": str(output)})
